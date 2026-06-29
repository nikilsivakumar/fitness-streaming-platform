import os


# Write a quick script to generate the docx using python-docx to perfectly mirror the js file content and layout
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()

# Base margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    
    # Headers/Footers
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("Fitness Streaming Analytics Platform — Master Handbook")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8)
    hrun.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun1 = fp.add_run("Page ")
    frun1.font.name = 'Calibri'
    frun1.font.size = Pt(8)
    frun1.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    # Add page number XML field
    f_xml = OxmlElement('w:fldSimple')
    f_xml.set(qn('w:instr'), 'PAGE')
    fp._p.append(f_xml)

# Palette
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT2 = "EEF3FA"
GREY = RGBColor(0x59, 0x59, 0x59)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '2')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'B9C6D9')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_heading_1(text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = NAVY
    
    # Add bottom border XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BLUE
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x31, 0x85, 0x9C)
    return p

def add_p(text, bold=False, italics=False, color=DARK, after=160):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after / 20)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italics = italics
    run.font.color.rgb = color
    return p

def add_quote_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(20)
    p.paragraph_format.left_indent = Inches(0.25)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italics = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.bold = True
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_table(widths, headers, rows):
    # Widths in dxa to inches conversion: 1 inch = 1440 dxa
    t = doc.add_table(rows=len(rows) + 1, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_cells = t.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_shading(hdr_cells[i], '1F3864')
        set_cell_borders(hdr_cells[i])
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = t.rows[r_idx + 1].cells
        shade = (r_idx % 2 == 1)
        for c_idx, cell_data in enumerate(row_data):
            set_cell_borders(row_cells[c_idx])
            if shade:
                set_cell_shading(row_cells[c_idx], LIGHT2)
                
            if isinstance(cell_data, list):
                p = row_cells[c_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(cell_data[0])
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                for line in cell_data[1:]:
                    np = row_cells[c_idx].add_paragraph()
                    np.paragraph_format.space_after = Pt(3)
                    run = np.add_run(line)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10.5)
            else:
                p = row_cells[c_idx].paragraphs[0]
                run = p.add_run(str(cell_data))
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                
    # Apply widths
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w / 1440.0)

def add_qa(q, a):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run("Q: " + q)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    r2 = p2.add_run("A: " + a)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

# ================= TITLE PAGE =================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("FITNESS STREAMING ANALYTICS PLATFORM")
r.font.name = 'Calibri'
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("The Master Engineering Handbook")
r.font.name = 'Calibri'
r.font.size = Pt(16)
r.font.italics = True
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
r = p.add_run("A Comparative Bronze → Silver → Gold Medallion Pipeline")
r.font.name = 'Calibri'
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Built Independently on Local, AWS, and Databricks")
r.font.name = 'Calibri'
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Nikil Sivakumar")
r.font.name = 'Calibri'
r.font.size = Pt(13)
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Data Analyst → Data Engineer | Chennai")
r.font.name = 'Calibri'
r.font.size = Pt(11)
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
r = p.add_run("github.com/nikilsivakumar/fitness-streaming-platform")
r.font.name = 'Calibri'
r.font.size = Pt(10)
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("Final portfolio reference document — architecture, build process, and interview material")
r.font.name = 'Calibri'
r.font.size = Pt(9)
r.font.italics = True
r.font.color.rgb = GREY

# ================= TOC =================
add_heading_1("Table of Contents", page_break=True)
p = doc.add_paragraph()
r = p.add_run("[Table of Contents Placeholder - Rendered in Word]")
r.font.italic = True

# ================= CH 1: OVERVIEW =================
add_heading_1("1. Project Overview & Thesis", page_break=True)
add_p("This project is a single pipeline built three times. The same Bronze → Silver → Gold medallion logic — the same validation rules, the same quarantine behavior, the same physiologically-grounded scoring formulas — is implemented independently on three infrastructure stacks: a hand-rolled Local stack (PySpark/DuckDB), AWS managed services (Kinesis/Lambda/Glue/Redshift/dbt), and Databricks (Delta Lake/Unity Catalog/Auto Loader/Workflows).")
add_quote_box("The comparison is the point. Anyone can demonstrate a pipeline on one stack. Being able to say \"here is the same logic, and here is why each platform makes that decision differently\" is the actual signal an interview is testing for — tool familiarity is replaceable; transferable engineering judgment is not.")

add_heading_2("What the pipeline does")
add_p("Five streaming record types from wearable/fitness sources — user_profile, wearable_event, workout_log, sleep_log, and nutrition_snapshot — flow through Bronze (raw, untransformed, append-only), Silver (validated, typed, deduplicated, quarantine-on-failure), and Gold (business aggregates: fatigue/recovery scoring, workout consistency, community cohort analytics, enriched user profiles).")

add_heading_2("Who this is for")
add_p("Built by Nikil, a data analyst in Chennai transitioning into data engineering, as the final portfolio project before interviews at roughly a 3-years-experience target level. The project is also a deliberate proving ground for a longer-term thesis: that physical training, financial independence, and operational discipline are the same underlying skill applied to different domains — the \"verify against real output, never trust a green checkmark\" discipline that runs through every phase of this build is the same discipline this project's author applies to coaching, to investing, and to the eventual self-sufficient farm-and-fitness system that sits beyond the 2026–2031 horizon. It is recorded here once, deliberately, rather than repeated throughout — the engineering chapters that follow speak for themselves.")

add_heading_2("Why three stacks, not one")
why3 = [
  ["Local (PySpark / DuckDB / StreamBus)", "Proves the pipeline logic itself, cheaply and fast-iterating, before paying for or depending on any managed service. Capped deliberately at Phase 4 — see §1.10 below."],
  ["AWS (Kinesis / Lambda / Glue / Redshift / dbt)", "Demonstrates fluency with managed cloud primitives — the lowest-operational-overhead way most companies actually run pipelines today."],
  ["Databricks (Delta Lake / Unity Catalog / Auto Loader / Workflows)", "Demonstrates the unified-lakehouse alternative — built second and explicitly as a structured migration exercise off the already-working AWS pipeline, since learning a new platform by translating something already understood is faster than learning it cold."],
]
add_table([3400, 5960], ["Stack", "What it proves"], why3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("Why the Local stack is capped at Phase 4")
add_p("Local's job was to prove the medallion logic once, on the cheapest iteration loop, before porting it to the two stacks that actually map to real hiring conversations. No real-world deployment of this pipeline would run as \"local PySpark + DuckDB\" in production — it would be AWS or Databricks. Continuing to build dbt-on-DuckDB, a Streamlit dashboard, and local orchestration would have been polishing a throwaway prototype instead of investing further time in the two stacks with real comparative payoff. This is a scope decision, not an unfinished phase — full reasoning in decisions.md §1.10.")

# ================= CH 2: ARCHITECTURE COMPARISON =================
add_heading_1("2. Architecture Comparison — All Three Stacks, Layer by Layer", page_break=True)
add_p("The table below is the single most interview-useful artifact in this project: the same row, three different answers, each defensible on its own terms.")

archRows = [
  ["Streaming transport", "Custom StreamBus (in-process, Kafka-equivalent abstraction — built after Docker/Kafka/Redpanda hit WSL2 incompatibilities on Windows 10)", "Kinesis Data Streams, 1 shard, deleted after each test session", "Auto Loader (cloudFiles), trigger(availableNow=True) — batch-triggered, not continuous"],
  ["Ingestion compute", "PySpark, run from CMD", "Lambda fitness-bronze-writer (Python 3.13, AWSSDKPandas layer)", "Databricks notebook (Serverless), live producer writing into a Unity Catalog Volume"],
  ["Storage format", "Hive-partitioned Snappy Parquet on local disk", "Hive-partitioned Parquet on S3 (fitness-streaming-nikil-2026)", "Delta Lake tables under Unity Catalog (ACID, time travel, schema evolution)"],
  ["Bronze→Silver transform", "bronze_to_silver.py, validation_rules.py (dispatch dict pattern)", "Glue ETL job fitness-bronze-to-silver", "5 independent notebooks, one per record type — user_profile via MERGE INTO, the other 4 via append"],
  ["Silver→Gold transform", "silver_to_gold.py", "Glue ETL job fitness-silver-to-gold", "4 independent notebooks — overwrite-on-recompute"],
  ["Query / validation layer", "DuckDB", "Athena (query-only, never moves pipeline data)", "Notebook SQL directly against Delta tables"],
  ["Warehouse / serving layer", "— (not built; scope-capped)", "Redshift Serverless — star schema, true SCD Type 2 on dim_user, loaded via Data API + dbt", "None needed — Unity Catalog + Delta already behave as the warehouse layer"],
  ["Catalog / lineage", "—", "Glue Data Catalog — lineage reconstructed manually from job definitions/S3 paths", "Unity Catalog — fully automatic lineage (11 upstream / 13 downstream tables surfaced with zero extra code)"],
  ["Orchestration", "Sequential script (not built; scope-capped)", "Glue job triggers, run manually per session", "Databricks Workflows — single Job, 11 tasks: producer → ingestion → 5 parallel Silver → barrier → 4 parallel Gold"],
  ["Transformation framework", "dbt (DuckDB profile — not built; scope-capped)", "dbt (Redshift profile) — fully built, 7 models, 33 tests, all green", "dbt (Databricks profile) — deferred, folds into the AWS-vs-Databricks comparison"],
  ["Cost discipline", "Zero marginal cost (local compute only)", "Kinesis shard deleted after use; Redshift Serverless auto-idles; Glue billed per-run, not continuous", "Free Edition (Serverless-only, quota-limited); producer runs bounded/finite, never continuous"],
]
add_table([2200, 2680, 2280, 2200], ["Layer", "Local", "AWS", "Databricks"], [[r[0], [r[1]], [r[2]], [r[3]]] for r in archRows])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("The one structural difference worth leading with in an interview")
add_p("On AWS, Gold lives in S3 as Parquet and has to be COPY'd into a genuinely separate system (Redshift) before it behaves like a queryable warehouse — two physical copies of the data exist, and they can drift out of sync if a COPY is missed. On Databricks, Unity Catalog + Delta Lake already provide ACID transactions, governance, and SQL-queryable access directly on the same files the Spark pipeline writes — there is no second system to load into. This single fact explains why AWS needed a whole additional phase (Phase 6: Redshift) that Databricks's equivalent track never required.")

add_heading_2("Repeated bug class, three different fixes")
add_p("The same underlying problem — a dimension table's first-arriving row was a CDC stream where one user could legitimately appear more than once, and downstream aggregates assumed one row per key — surfaced three separate times across this project, each time requiring a fix specific to the platform it was found on:")
add_bullet("Local & AWS Glue (gold_user_profile_enriched): no dedup on a flat select over a CDC-style user_profile source. Fixed with a row_number() window keyed on latest timestamp.")
add_bullet("Databricks (silver.user_profile): same root issue, fixed upstream and more cleanly — MERGE INTO at Silver guarantees exactly one row per user before Gold ever reads it, so no Gold-layer dedup is needed at all.")
add_bullet("AWS Redshift dbt (dim_user.effective_date): a related but distinct manifestation — the dimension's first SCD2 version was stamped with the ETL's load timestamp rather than an open-ended business-validity date, causing every same-day fact row to silently fail the SCD2 join. Fixed in stg_dim_user with an effective_date_for_join column.")
add_quote_box("Naming this as one story — \"the same architectural gap, found three times, fixed three different ways depending on what was actually available downstream\" — is stronger than three unrelated bug-fix anecdotes.")

# ================= CH 3: STEP BY STEP BUILD PROCESS =================
add_heading_1("3. Step-by-Step Build Process", page_break=True)
add_p("This chapter is the condensed build log — phase by phase, what was built, what broke, and how it was verified. Full reasoning for every decision referenced here lives in decisions.md; this chapter is the narrative spine.")

add_heading_2("3.1 Local track (Phases 1–4)")
localPhases = [
  ["Phase 1", "StreamBus (in-process broker), wearable simulator (5 record types, weighted distribution, 5% injected bad data), Bronze writer (Hive-partitioned Snappy Parquet)."],
  ["Phase 2", "Five reference CSVs generated: user_profiles (50 users), gym_master (5 gyms), workout_catalog (MET values), medical_conditions (risk modifiers), date_dimension (365 rows)."],
  ["Phase 3", "Bronze → Silver: validation_rules.py (dispatch-dict pattern, physiologically grounded thresholds — HR 30–220bpm, HRV>0, sleep 0–16hrs, RPE 1–10), bronze_to_silver.py (dedup, type-cast with sentinel fill, validate, split valid/quarantine, enrich via reference joins). Hit and fixed a JAVA_GATEWAY_EXITED error (Java/Hadoop misconfigured on Windows) independently."],
  ["Phase 4", "Silver → Gold: fatigue score, recovery readiness, overtraining flag, workout consistency, community analytics — re-validated against the real Silver schema rather than reused from an earlier unvalidated draft."],
]
add_table([1400, 7960], ["Phase", "What was built"], localPhases)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
add_p("Phases 7–9 (dbt/DuckDB, Streamlit, local orchestration) were deliberately not built — see §1.10 / decisions.md §1.10.")

add_heading_2("3.2 AWS track (Phases 5–7)")
awsPhases = [
  ["Phase 5", "S3 bucket (fitness-streaming-nikil-2026), IAM reused from a prior retail project (retail-pipeline-dev, deliberately lacking iam:CreateRole and ec2:AuthorizeSecurityGroupIngress). Kinesis → Lambda → S3 Bronze built, tested with 100 records across 4 batches, verified, then torn down per cost discipline."],
  ["Phase 5 (Glue)", "Glue ETL jobs fitness-bronze-to-silver and fitness-silver-to-gold ported faithfully from Local. Both jobs verified SUCCEEDED via CLI JobRunId poll, not just trusted. Glue Crawler + Athena catalog fitness_streaming_db built; all 4 Gold tables crawled and queryable."],
  ["Phase 6", "Redshift Serverless (fitness-namespace / fitness-workgroup, 8 RPU). IAM role RedshiftS3ReadRole created via Console since retail-pipeline-dev correctly cannot create IAM roles. Star schema loaded from S3 Gold via COPY, including a SUPER-type nested column (muscle_group_breakdown) loaded with SERIALIZETOJSON. SCD Type 2 schema scaffold added to dim_user."],
  ["Phase 7", "dbt (Redshift profile). Discovered the real schema lives in fitness_dw, not the default dev database, via macro introspection. Built staging/marts/tests/macros. Found and fixed the dim_user.effective_date SCD2 bug (see §2 above) via a manual row-count check after a 100%-green test suite missed it entirely. Exported the live schema to version_aws/redshift/schema.sql, closing a previously-flagged gap."],
]
add_table([1400, 7960], ["Phase", "What was built"], awsPhases)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("3.3 Databricks track (DB-1 through DB-5)")
dbPhases = [
  ["DB-1", "Unity Catalog: catalog fitness_streaming, schemas bronze/silver/gold/reference, Volume fitness_streaming.bronze.raw_events. Learned that a Volume is a registered Unity Catalog object — os.makedirs() can create subfolders inside one, but not register a new top-level Volume."],
  ["DB-2", "Live producer (reuses Local's generate_*() functions unchanged) writing into the Volume; Auto Loader (5 independent cloudFiles streams, trigger(availableNow=True)) ingesting into Bronze Delta tables. Verified 450/450 rows landed with zero loss/duplication — exact match to producer output."],
  ["DB-3", "Bronze → Silver, one notebook per record type. user_profile uses MERGE INTO (idempotency proven by a second no-op run); the other four use append + dedup-on-(user_id, timestamp). wearable_event's quarantine table confirmed non-empty, matching the intentional 5% bad-data injection."],
  ["DB-4", "Silver → Gold, all 4 tables. Found and fixed the same reference-enrichment gap three separate times (medical_risk_modifier/gym_type missing from user_profile; statement-scoped MERGE WITH SCHEMA EVOLUTION needed since Serverless rejects session-wide spark.conf.set; muscle_groups array serialized as Spark's non-JSON bracket format, needing regexp_replace + split instead of Local's quote-swap + from_json). Closed the gold_community_analytics null-cohort open item with .fillna(\"Unprofiled\")."],
  ["DB-5", "Databricks Workflows — 11-task DAG (producer → ingestion → 5 parallel Silver → coarse barrier → 4 parallel Gold), retries split by failure class (2 for transient infra, 0 for deterministic logic). Caught and removed a duplicate run_producer() call before the first run that would have silently doubled data volume on every execution. First full run: SUCCEEDED, 5m19s, all 11 tasks green, verified via real row-count output."],
]
add_table([1100, 8260], ["Stage", "What was built"], dbPhases)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("3.4 The single strongest verification story in the project")
add_p("During AWS Phase 7, dbt run reported 7/7 models successful and dbt test reported 31/31 tests passing — a completely clean run. A manual row-count check anyway (10 raw fact rows vs. 0 mart rows) revealed mart_user_health_score was entirely empty. The root cause — every dim_user row carrying an identical, literal load timestamp as its effective_date, so even same-day fact rows failed the SCD2 join — was invisible to every test that had been written, because no test had been written to check for it. A permanent regression test (assert_health_score_no_dropped_rows) was added afterward specifically so this exact failure mode would be caught automatically in the future.")
add_quote_box("A green test suite is necessary, not sufficient. This is the headline interview story for the whole project: the discipline of checking real output caught a 100%-empty mart table that 31 passing tests did not.")

# ================= CH 4: EXTENDED ARCHITECTURE =================
add_heading_1("4. Extended Architecture — Endpoints Beyond Redshift and Delta", page_break=True)
add_p("Redshift Serverless (AWS) and Unity Catalog Delta tables (Databricks) are this project's only built endpoints. Neither is the only possible one. This chapter documents the realistic alternatives that were reasoned through but deliberately not built — because knowing what else could sit downstream, where it would physically live, and when you would not reach for it, is itself the kind of judgment this project exists to demonstrate. Full source reasoning: decisions.md §6.")

add_heading_2("4.1 AWS side — Snowflake as an alternative to Redshift, not an extension of it")
add_p("Snowflake would sit completely outside the AWS account — a separate vendor account, separate billing, separate IAM/RBAC, with no shared identity with account 811821010882. The only connective tissue is a storage integration: an IAM role in the AWS account that Snowflake is granted permission to assume, so its own COPY INTO can read Parquet directly from the S3 Gold prefix. No AWS compute is involved in that read at all — Snowflake's own \"virtual warehouse\" compute does the work entirely outside the AWS boundary.")
snowWhen = [
  ["Multi-cloud requirement, or BI/finance already standardized on Snowflake", "Very common in practice — the data engineering team builds on AWS/Databricks while the rest of the org already pays for Snowflake for unrelated historical reasons."],
  ["Independent compute scaling, not just storage scaling", "Snowflake's separate virtual warehouses let a heavy BI workload and a heavy batch-transform workload run on two pools that never contend — Redshift Serverless is a single compute pool per workgroup."],
  ["Lower operational overhead", "No distribution keys, sort keys, VACUUM/ANALYZE, or security-group/publiclyAccessible networking — the exact friction this project hit directly needing to open a security group for every local dbt session (decisions.md §3.10)."],
  ["Cross-org data sharing", "Snowflake Secure Data Sharing if Gold data ever needed to reach an external partner without copying it."],
]
add_table([3200, 6160], ["When Snowflake is the right call", "Why"], snowWhen)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
add_p("Redshift remains the right call here specifically because the whole pipeline already lives in one AWS account end to end — IAM, S3, Glue, and Redshift share one identity and billing boundary, and the Data API authenticating directly off retail-pipeline-dev's IAM identity (no separate password store) is a genuine advantage that a second vendor account would dilute for no corresponding gain at this project's scale.")

add_heading_2("4.2 AWS side — BI/visualization layer, and where each tool actually lives")
biRows = [
  ["Amazon QuickSight", "Inside the AWS account — a managed AWS service on the same bill, same region", "Native Redshift/Athena/S3 connector, IAM-based access — no separate credential store", "Lowest friction precisely because it never leaves the AWS network boundary"],
  ["Tableau", "Separate vendor entity by default (Tableau Cloud = pure SaaS, zero AWS footprint); Tableau Server can be self-hosted on EC2 inside the account but is still a separately licensed product", "JDBC/ODBC Redshift driver, live connection or scheduled extract", "Hits the same \"Redshift Serverless not publicly accessible by default\" friction as local dbt sessions — needs a permanent (not dev-IP-scoped) security group rule or a self-hosted Bridge/Server"],
  ["Power BI", "Microsoft SaaS, entirely outside the AWS account", "Power BI's Redshift connector + an on-prem data gateway if Redshift stays private", "Same security-group friction as Tableau; gateway pattern mirrors Tableau Bridge"],
]
add_table([1500, 2700, 2580, 2580], ["Tool", "Where it lives", "How it connects", "Operational note"], biRows)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("4.3 AWS side — when this project would actually use SageMaker")
add_p("SageMaker is not a warehouse or BI replacement — it is a parallel branch reading from Silver/Gold, not a further step after Redshift. It would sit inside the AWS account, training against S3 Parquet or via Athena/Redshift Data API, and serving predictions through a real-time endpoint or a batch transform job.")
add_bullet("the current overtraining_flag is a same-day, reactive, deterministic rule. A model trained on rolling 7–14 day HRV trend, sleep-debt trajectory, and training load could predict risk before the threshold is crossed, by learning non-linear interactions the additive formula can't capture. This is the single most defensible \"ML adds real value here\" answer for this dataset.", bold_prefix="Predictive overtraining risk: ")
add_bullet("declining workout_consistency is a leading indicator for a client about to drop a coaching program — ties the technical pipeline directly to the actual personal-training business this data models.", bold_prefix="PT-client churn risk: ")
add_bullet("clustering on fatigue/recovery/consistency feature vectors could surface cohorts beyond the two explicit dimensions (job_type, fitness_goal) already in Gold.", bold_prefix="Unsupervised cohort discovery: ")
add_p("Cost discipline note: a deployed real-time SageMaker endpoint bills continuously while running — the same always-on category this project has avoided everywhere else (Kinesis deleted after use, Redshift auto-idling, Glue/Lambda triggered-not-continuous). A training job or batch transform fits the existing discipline; a permanently-deployed endpoint would need the same cost conversation already had about EC2.")
add_quote_box("Why it isn't built: the current fatigue/recovery scoring is interpretable and explicitly grounded in cited physiological research. In a health/wellness context, that interpretability has real value — a user or coach can be told exactly why a score is what it is. Recognizing that not every scoring problem benefits from ML, and being able to say specifically why for this dataset, is a more senior answer than defaulting to \"and then I'd add a model.\"")

add_heading_2("4.4 Databricks side — building a dashboard directly off the Gold Delta endpoint")
dbDash = [
  ["Databricks SQL + Lakeview dashboards", "Most native option. Built on a serverless SQL Warehouse, queries Gold Delta tables directly through the same Unity Catalog permissions already governing everything else, zero data movement, inherits automatic lineage for free. The natural next step (DB-6)."],
  ["Databricks Apps (Streamlit, natively hosted)", "Lets the Streamlit dashboard originally scoped for Local's abandoned Phase 8 be reused rather than rebuilt — the same codebase pointed at DuckDB locally and at a Databricks SQL Warehouse here, with only the connection string differing."],
  ["Power BI / Tableau via Partner Connect", "Databricks's direct equivalent of \"Tableau connects to Redshift\" — same external-vendor-account caveat, but the connection is HTTPS/token-based rather than raw Postgres-over-a-security-group, meaningfully less networking friction than the Redshift-side experience."],
]
add_table([2900, 6460], ["Option", "Why / how"], dbDash)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("4.5 Databricks side — can the endpoint become Snowflake too?")
add_p("Yes, via three genuinely different mechanisms:")
add_bullet("Databricks's open, vendor-agnostic protocol for sharing live Delta tables across organizations without copying data. Snowflake has native Delta Sharing support — Gold Delta tables could be shared directly into a Snowflake account, governed entirely from Unity Catalog, with no ETL hop and no format conversion. The cleanest mechanism, and the one worth leading with, since it's an open standard rather than a Databricks-only trick.", bold_prefix="Delta Sharing: ")
add_bullet("the reverse direction — Databricks queries a live Snowflake table without copying it in, useful when Snowflake is the system of record for some other part of the org.", bold_prefix="Unity Catalog Lakehouse Federation: ")
add_bullet("write Gold as Parquet to cloud storage and let Snowflake's own COPY INTO ingest it — structurally identical to the AWS-side Snowflake option in §4.1, just sourced from Databricks instead of Glue.", bold_prefix="Plain export/COPY: ")
add_p("This matters as a real-world pattern, not a hypothetical: it's extremely common for a data/ML team to standardize on Databricks while the rest of an org's BI consumers already run on Snowflake for unrelated historical reasons. Delta Sharing exists specifically to make that split organizationally survivable without forcing either team to migrate.")

# ================= CH 5: INTERVIEW SCENARIOS =================
add_heading_1("5. Interview Scenarios", page_break=True)
add_p("A curated set of the highest-value questions this project should be able to answer under pressure. The full 50+ entry bank lives in docs/interview_qa.md — this is the distilled version.")

add_heading_2("5.1 Framing & motivation")
add_qa("Why build the same pipeline three times instead of going deep on one stack?",
       "Most candidates can show one pipeline on one stack. Building the same Bronze→Silver→Gold logic on a hand-rolled local stack, on AWS managed services, and on Databricks lets me talk about why each platform makes a given decision differently, not just that I used the platform. It's a comparison engine, not three disconnected demos.")
add_qa("What's the single biggest differentiator versus a typical bootcamp pipeline project?",
       "The decision log behind it. Anyone can follow a tutorial. Being able to say \"Silver uses MERGE INTO because rows have stable identity, Gold uses full overwrite because aggregates don't — and here's why that's true in PySpark, in Glue, and in Databricks\" shows the reasoning transfers across tools.")

add_heading_2("5.2 Architecture trade-offs")
add_qa("Why does Silver use MERGE INTO and Gold use full overwrite? Isn't that inconsistent?",
       "It looks inconsistent until you look at the grain. Silver rows have stable identity — a user, a workout session — so there's a real merge key. Gold rows are aggregates with no natural row-level key. Inventing a synthetic key just to use MERGE INTO would be solving a problem that doesn't exist.")
add_qa("Why Kinesis/Lambda/Glue on AWS instead of EC2 or EMR?",
       "EC2/EMR would make the AWS leg look like \"Local stack on bigger hardware,\" which loses the comparative point of the project. Kinesis/Lambda/Glue represents the managed-primitives, low-ops posture AWS is actually good at, contrasting meaningfully with both Local (self-managed) and Databricks (managed platform).")
add_qa("Why didn't you put a data warehouse in front of the Databricks Gold tables, the way you did with Redshift on AWS?",
       "Because Databricks doesn't need that hop. Unity Catalog + Delta Lake already provide ACID transactions, governance, and SQL-queryable access on the same files the pipeline writes — there's no second system to load into. A Databricks SQL Warehouse is compute, not a second copy of the data. Needing Redshift on AWS and not needing an equivalent on Databricks is itself one of the cleanest structural contrasts in the whole project.")

add_heading_2("5.3 The extended-architecture questions (likely follow-ups once the build is presented)")
add_qa("If Redshift turned out to be the wrong long-term choice, what would you actually do?",
       "Snowflake via a storage integration reading Parquet directly from the existing S3 Gold prefix — no Lambda or Glue involved in that read. I'd reach for it specifically if BI consumers already standardized on Snowflake elsewhere in the org, or if I needed independently-scaling compute pools rather than Redshift Serverless's single workgroup pool. I wouldn't reach for it just because it's popular — staying in one AWS account with one IAM identity has real, current value for this project.")
add_qa("How would you put a BI tool in front of this without breaking your cost/networking discipline?",
       "QuickSight first, because it's a managed AWS service on the same bill and never leaves the AWS network boundary — no security-group exposure at all, unlike Tableau or Power BI, which would hit the exact same \"Redshift Serverless isn't publicly accessible by default\" friction I already documented for local dbt sessions.")
add_qa("Where would you actually use SageMaker on this project, and where would you deliberately not?",
       "Predicting overtraining risk before the threshold is crossed, using rolling HRV/sleep/training-load trends — that's a place a model could learn interactions a hand-tuned additive formula can't. I would not replace the existing fatigue/recovery scoring outright, because it's interpretable and grounded in cited research; in a health context that explainability is worth more than an unproven accuracy gain.")
add_qa("Can the Databricks Gold tables feed Snowflake without rebuilding the pipeline?",
       "Yes — Delta Sharing, an open protocol Snowflake natively supports, lets me share the live Delta tables into a Snowflake account with no copy and no format conversion, governed from Unity Catalog. That's the real answer to \"what if the rest of the org is on Snowflake\" without forcing a platform migration.")

add_heading_2("5.4 The hardest question: \"What would you do differently?\"")
add_p("A strong, specific answer rather than a vague one: \"I'd write the dim_user.effective_date regression test (assert_health_score_no_dropped_rows) before the bug, not after — and more generally, I'd treat 'no test exists for this failure mode yet' as its own category of risk, separate from 'all existing tests pass.' I'd also port the Databricks reference-enrichment gap fix back to AWS proactively once I found it the second time on Databricks, rather than waiting to discover whether the same gap existed on the AWS Glue side too.\"")

# ================= CH 6: SCALING SCENARIOS =================
add_heading_1("6. Scaling Scenarios", page_break=True)
add_p("How each stack's current design would (or would not) hold up under materially higher load — a standard senior-level interview probe.")

add_heading_2("6.1 10x more users / wearable events per second")
scale10 = [
  ["Local", "Breaks down structurally — single-process StreamBus and PySpark on a laptop have no horizontal scaling story. This is exactly why Local is capped at Phase 4 and never positioned as a deployment target."],
  ["AWS", "Increase Kinesis shard count (each shard handles ~1MB/s in or 1,000 records/s); Lambda scales concurrency automatically per shard via the event source mapping, up to account concurrency limits; Glue ETL scales by increasing DPUs/worker count. The architecture doesn't change shape — it adds capacity at each existing layer."],
  ["Databricks", "Auto Loader scales by increasing the SQL/cluster size behind the streams or splitting record types across more parallel streams; trigger(availableNow=True) batch-style ingestion would likely need to move toward continuous triggering at high enough velocity, since waiting for \"available now\" batches stops being efficient once arrival is truly continuous."],
]
add_table([1500, 7860], ["Stack", "What changes"], scale10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("6.2 10x more historical data in Gold / the warehouse")
add_bullet("AWS: Redshift Serverless scales RPU allocation automatically within configured limits; sort/distribution key choices on the fact tables start to matter more at this volume — currently not tuned, since data volume has been small (10s of rows). Worth naming as a real, currently-deferred optimization rather than claiming it's already handled.")
add_bullet("Databricks: Delta's file-compaction (OPTIMIZE, Z-ORDER) becomes relevant for query performance at scale — not needed yet at current volumes, but a known, named lever rather than an unknown one.")

add_heading_2("6.3 Real-time SLA requirements (sub-second freshness)")
add_p("Both AWS and Databricks's current designs are intentionally batch-triggered (Glue runs per session; Auto Loader uses availableNow=True), which is correct for this project's actual requirements but would need to change under a true real-time SLA: on AWS, Lambda's per-record processing already is near-real-time at the Bronze layer, but Glue's Silver/Gold steps would need to move toward Kinesis Data Analytics / Spark Structured Streaming with continuous triggers instead of scheduled batch runs. On Databricks, Auto Loader would drop availableNow=True for continuous triggering, and Gold aggregates would need to move from full-recompute-on-batch toward incremental/streaming aggregation — a materially different engineering problem, not just a config flag.")

add_heading_2("6.4 Multiple teams / multi-tenant access")
add_p("This is where Databricks's Unity Catalog and AWS's Glue Data Catalog + IAM diverge most under scale: Unity Catalog gives column- and row-level access control plus automatic lineage natively across catalogs/schemas, which is exactly what a multi-team org needs without bolting on a separate governance tool. On AWS, equivalent fine-grained governance would mean adding Lake Formation on top of Glue/S3/Redshift — a real additional component, not something Glue alone provides.")

# ================= CH 7: FAILURE SCENARIOS =================
add_heading_1("7. Failure Scenarios", page_break=True)
add_p("What happens — and what the recovery path is — when each layer fails. This project's quarantine-not-drop and Bronze-immutability principles make most of these answers concrete rather than hypothetical.")

failRows = [
  ["Kinesis shard throttled / records rejected", "AWS", "Lambda's event source mapping retries automatically; Kinesis itself retains records for 7 days, so a transient Lambda failure doesn't lose data — it delays processing. Recovery: investigate Lambda concurrency/throttle limits, no data loss to recover from."],
  ["Glue job fails mid-run", "AWS", "Bronze is immutable and untouched by a failed Glue job — fix the bug, delete the partial/bad Silver partition if one was written, and rerun from Bronze. Nothing upstream needs to be redone."],
  ["Redshift COPY fails (e.g., the SUPER/SERIALIZETOJSON or type-width bugs hit this session)", "AWS", "Gold Parquet remains safe and untouched in S3 regardless of COPY outcome — fix the DDL/COPY options and rerun COPY. The failure is fully recoverable because the source of truth (S3) was never at risk."],
  ["dbt test suite is 100% green but the mart is wrong", "AWS", "This actually happened (§3.4). Recovery path: never treat a green test suite as sufficient — run a manual row-count/spot-check against real numbers before declaring a model done, and write a permanent regression test for the exact failure mode once found."],
  ["A Databricks Job task fails (Silver or Gold)", "Databricks", "Retries=0 by design for Silver/Gold — a deterministic logic failure should alert immediately, not retry and delay the alert. Recovery: fix the notebook logic, rerun the Job (or just the failed task) — Bronze/Silver upstream of the failure point is untouched, same Bronze-immutability principle as AWS."],
  ["A Databricks Job task fails (producer or ingestion)", "Databricks", "Retries=2 by design — these are typically transient infra issues (package install hiccups, cluster init) that may genuinely succeed on a second attempt, unlike a deterministic logic bug."],
  ["Notebook state corruption (AMBIGUOUS_REFERENCE-style errors that survive a restart)", "Databricks", "A notebook's .py file and its live Python process state are two different things — out-of-order cell execution can leave stale dataframe lineage behind. Recovery: hard cluster Detach & Re-attach, then Run All from the top, never partial/out-of-order execution while debugging."],
  ["Bad/malformed record arrives at Silver validation", "All three", "Never silently dropped, never fails the whole batch — quarantined with a failure_reason, verified by confirming quarantine counts match intentionally-injected bad data exactly (wearable_event's 5% injection, confirmed non-empty on every stack)."],
  ["A dimension's data arrives after the facts that depend on it (late-arriving dimension)", "AWS & Databricks", "Two different correct answers depending on context: Databricks's gold_community_analytics explicitly labels unmatched users \"Unprofiled\" rather than dropping or nulling them, preserving their activity data. AWS's SCD2 fix instead corrected the dimension's effective_date itself, because in that case the late arrival was an ETL-timestamp artifact, not a genuine business late-arrival — recognizing which of the two situations you're in is the actual skill, not applying one fix everywhere."],
]
add_table([2700, 1100, 5560], ["Failure", "Stack", "What happens / recovery path"], failRows)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)

add_heading_2("7.1 The recovery principle underneath all of these")
add_quote_box("Every layer is a checkpoint, not a single point of failure. Bronze is immutable, so any downstream failure can always be recovered by fixing logic and reprocessing from Bronze — never from re-ingesting source data. This is true identically on AWS and Databricks, and is the reason \"what happens if X fails\" almost always has a calm, specific answer rather than \"I'd have to start over.\"")

# ================= CH 8: STANDING PRINCIPLES =================
add_heading_1("8. Standing Engineering Principles", page_break=True)
add_p("The disciplines that were applied consistently across all three stacks, all ten-plus build sessions, and both shared-documentation passes. These are the answer to \"how do you work,\" independent of any specific tool.")
add_numbered("Verify against real output before declaring anything done. Every real bug in this project — across all three stacks — was caught by inspecting actual run output (schema, row counts, specific values), never by re-reading code or trusting a green checkmark.")
add_numbered("Bugfixes are committed separately from feature work, every time, no exceptions.")
add_numbered("Logic stays identical across stacks. Porting a transformation means referencing the existing implementation, not redesigning it for the new platform — the comparison only means something if the underlying logic is genuinely the same.")
add_numbered("Blast-radius isolation over DRY. One notebook/script per record type, so a bug in one can't cascade into the other four — an explicit, named trade against code reuse.")
add_numbered("Quarantine, never silently drop or hard-fail, on bad data.")
add_numbered("Cost discipline as an operational habit, not an afterthought. Kinesis shards deleted after use, Redshift Serverless auto-idling understood and trusted (not manually \"paused\" the way a provisioned cluster would need), Databricks ingestion batch-triggered rather than continuous — each one a deliberate choice tied to actual billing mechanics, checked rather than assumed.")
add_numbered("Investigate suspected leftovers before deleting anything. The Redshift staging tables looked like clutter and turned out to be load infrastructure — verified with a real comparison query before any decision was made.")
add_numbered("Document decisions with the alternatives that were rejected, not just the choice made. A decision without a rejected alternative reads as the only option considered, which is rarely true and never as convincing.")
add_numbered("IAM/security boundaries hold under real pressure. retail-pipeline-dev correctly got blocked from iam:CreateRole and ec2:AuthorizeSecurityGroupIngress multiple times across this project, each time requiring the higher-privileged console path instead — proof the least-privilege design was real, not theoretical.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("End of handbook.")
r.font.name = 'Calibri'
r.font.size = Pt(11)
r.font.italics = True
r.font.color.rgb = GREY

doc.save("Fitness_Streaming_Platform_Master_Handbook.docx")
print("Successfully generated Word document!")